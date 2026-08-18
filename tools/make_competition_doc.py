# -*- coding: utf-8 -*-
"""Generate the competition documentation (PDF + DOCX) for 牛来举铁."""
import hashlib
import os

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer, PageBreak,
    Table, TableStyle, NextPageTemplate, KeepTogether
)
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.graphics.shapes import Drawing, Rect, String, Line, Polygon


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(ROOT, "outputs", "pdf")
os.makedirs(OUT_DIR, exist_ok=True)

PDF_PATH = os.path.join(OUT_DIR, "牛来举铁-微信小程序赛事说明文档.pdf")
DOCX_PATH = os.path.join(OUT_DIR, "牛来举铁-微信小程序赛事说明文档.docx")

# ---------- Fonts ----------
FONT = "MSYH"
FONT_B = "MSYHBD"
pdfmetrics.registerFont(TTFont(FONT, r"C:\Windows\Fonts\msyh.ttc", subfontIndex=0))
pdfmetrics.registerFont(TTFont(FONT_B, r"C:\Windows\Fonts\msyhbd.ttc", subfontIndex=0))

NAVY = colors.HexColor("#0a1530")
PURPLE = colors.HexColor("#5645d4")
INK = colors.HexColor("#1a1a1a")
SLATE = colors.HexColor("#5d5b54")
HAIR = colors.HexColor("#e5e3df")
SURFACE = colors.HexColor("#f6f5f4")


def st(name, **kw):
    base = dict(fontName=FONT, fontSize=10.5, leading=17, textColor=INK,
                spaceAfter=6, alignment=TA_LEFT)
    base.update(kw)
    return ParagraphStyle(name, **base)


S_H1 = st("H1", fontName=FONT_B, fontSize=16, leading=22, textColor=NAVY,
          spaceBefore=18, spaceAfter=8, keepWithNext=1)
S_H2 = st("H2", fontName=FONT_B, fontSize=12.5, leading=17, textColor=PURPLE,
          spaceBefore=12, spaceAfter=6, keepWithNext=1)
S_H3 = st("H3", fontName=FONT_B, fontSize=11, leading=16, textColor=INK,
          spaceBefore=8, spaceAfter=4, keepWithNext=1)
S_P = st("Body")
S_BULLET = st("Bullet", leftIndent=16, bulletIndent=6, spaceAfter=3)
S_CAPTION = st("Caption", fontSize=9, leading=13, textColor=SLATE)
S_CELL = st("Cell", fontSize=8.6, leading=12.5, spaceAfter=0)
S_CELLB = st("CellB", fontName=FONT_B, fontSize=8.8, leading=12.5, textColor=NAVY, spaceAfter=0)
S_TITLE = st("Title", fontName=FONT_B, fontSize=26, leading=34, textColor=colors.white,
              alignment=TA_CENTER)
S_SUB = st("Sub", fontSize=13, leading=20, textColor=colors.HexColor("#a4a097"),
           alignment=TA_CENTER)
S_TOC1 = st("TOC1", fontName=FONT_B, fontSize=11, leading=20, textColor=INK)
S_TOC2 = st("TOC2", fontSize=10, leading=17, textColor=SLATE, leftIndent=16)


# ---------- Content ----------
CONTENT = [
    ("h1", "一、作品概述"),
    ("h2", "1.1 作品简介"),
    ("p", "「牛来举铁」是一款基于微信云开发的原生微信小程序，专注「健身记录」这一件小事。"
          "产品的第一性原理是三个动作连成的最短闭环：记一次训练 → 看见进步 → 分享成果。"
          "围绕这个闭环，小程序提供了首次建档引导、专业动作库训练记录、历史统计与个人最好成绩（PR）追踪、"
          "训练分享卡四大能力，让用户在 30 秒内完成一次训练记录，并在数据中看见自己的进步。"),
    ("p", "当前版本 v1.2.0，共 6 个页面、4 个数据库集合、6 个云函数，内置 59 个专业动作与 4 套男女差异化训练模板，"
          "动作库涵盖胸、背、腿、臀腿、肩、手臂、核心 7 大部位，全部动作带中英文术语、器械类型与推荐重量区间。"),
    ("h2", "1.2 核心亮点"),
    ("bullet", "最短闭环：建档 → 记录 → 统计 → 分享，主路径一步不多、一步不少"),
    ("bullet", "专业动作库：59 个动作按 7 大部位分类，中英文对照，标注器械，男女分别给出推荐重量区间"),
    ("bullet", "自定义动作：内置动作之外可随时添加自定义动作，满足不同训练需求"),
    ("bullet", "模板训练：4 套男女差异化模板（全身分化男/女版、PPL、上下肢分化），训练中可随时更换模板"),
    ("bullet", "动作间歇计时器：每组动作可设 1–30 分钟休息倒计时，底部常驻显示，到点震动提醒"),
    ("bullet", "本地草稿兜底：断网不丢数据，训练草稿本地保存，启动时自动补传云端"),
    ("bullet", "分享卡：Canvas 本地绘制，白卡/照片背景可选，一键保存相册或转发好友"),
    ("bullet", "Notion 风格设计语言：设计令牌集中管理，全 rpx + 安全区适配，真机无横向滚动"),
    ("h2", "1.3 产品定位与目标用户"),
    ("p", "产品定位为「轻量、专业、有反馈」的健身记录工具，面向健身新手与进阶爱好者："
          "新手需要引导与专业参照（练什么、多重、几组），进阶用户需要数据沉淀与进步反馈（PR、趋势），"
          "两类用户都希望训练成果能体面地分享到微信朋友圈。"),

    ("h1", "二、创作背景"),
    ("h2", "2.1 需求洞察"),
    ("p", "近年来国内健身人口持续增长，力量训练从「小众爱好」走向「大众生活方式」。"
          "但在实际使用中，健身记录存在明显断层：专业的健身 App 功能强大却学习成本高，"
          "备忘录与纸质记录无法沉淀数据，微信群打卡热闹但缺乏统计与反馈。"
          "用户真正需要的不是「更多功能」，而是一个能快速记下来、看得见进步、方便晒出去的工具。"),
    ("h2", "2.2 用户痛点"),
    ("bullet", "记录难坚持：手动记录繁琐，传统工具输入成本高，训练中打断节奏"),
    ("bullet", "缺乏专业参照：新手不知道动作名称、器械类型、该用多重、该做几组"),
    ("bullet", "数据不可见：记录散落在备忘录/Excel，无法形成趋势与个人最好成绩"),
    ("bullet", "分享无出口：想发朋友圈晒训练成果，需要手动拼图，门槛高"),
    ("h2", "2.3 竞品分析"),
    ("table", {
        "headers": ["同类产品", "优势", "不足", "我们的机会"],
        "widths": [24, 26, 26, 24],
        "rows": [
            ["专业健身 App（Keep 等）", "功能全、社区成熟", "功能臃肿，记录链路长，新手易迷失", "聚焦单点，30 秒完成记录"],
            ["健康管理 App（薄荷等）", "饮食与健康数据全", "训练记录弱，偏重摄入管理", "专注力量训练数据"],
            ["纸质 / 备忘录 / Excel", "零成本", "无统计、无反馈、易丢失", "数据自动沉淀 + 专业参照"],
            ["微信群文字打卡", "门槛低、有社交氛围", "无数据积累、无个人反馈", "数据 + 分享闭环"],
        ],
    }),
    ("h2", "2.4 机遇与意义"),
    ("p", "微信生态为健身记录提供了天然土壤：小程序免安装、随开随用；训练成果可一键生成分享卡进入朋友圈，"
          "形成「记录 → 反馈 → 分享 → 带动身边人」的正循环。技术层面，微信云开发提供云函数、云数据库、"
          "免运维托管，让个人开发者也能做出可靠的产品。这与「健康中国」全民健身的大背景高度契合。"),

    ("h1", "三、创作思路"),
    ("h2", "3.1 第一性原理：最短闭环"),
    ("p", "产品设计的第一性不是「做一堆页面」，而是三个动作连成的最短闭环：记一次训练 → 看见进步 → 分享成果。"
          "所有功能都围绕这条闭环展开，凡是不服务于闭环的能力都被裁剪或降级为可选增量。"),
    ("diagram", "loop"),
    ("p", "上一版按「功能清单」堆砌页面，导致登录与建档分离、模板与分享强依赖云端与画布、样式全局复用造成真机横向漂移。"
          "重建版以闭环为纲：数据「按需加载 + 本地兜底」，样式「盒子模型统一 + 移动端优先」，"
          "非核心功能保留但架构上可裁剪，不再压在主路径上。"),
    ("h2", "3.2 产品设计原则"),
    ("bullet", "最短闭环优先：主路径一步不多、一步不少，30 秒内完成一次记录"),
    ("bullet", "数据可靠优先：本地草稿 + 云端同步，断网不丢数据，启动自动补传"),
    ("bullet", "专业但不吓人：用标准数据做「推荐」而非「说教」，给出区间与参照"),
    ("bullet", "移动端优先：全 rpx、全局 border-box、安全区适配、禁止横向滚动"),
    ("h2", "3.3 核心功能设计"),
    ("h3", "首次建档引导"),
    ("p", "三步引导：协议勾选 → 录入性别、年龄、身高、体重、健身目标（每周频率可选）→ 展示体质指标结果（BMI、"
          "体脂参考区间、基础代谢 BMR）。建档即算指标，让用户第一次打开就知道自己的起点，训练页再根据性别与体重"
          "给出每个动作的推荐重量区间。"),
    ("h3", "训练记录"),
    ("p", "支持自由训练与模板训练两种模式。自由训练按部位选择动作；模板训练一键带入专业分化计划，训练中可随时换模板。"
          "每个动作按组录入次数与重量（自重动作重量为 0），每组动作可设置 1–30 分钟间歇倒计时，底部常驻显示总组数与总容量，"
          "完成训练后进入分享页。"),
    ("h3", "历史统计"),
    ("p", "历史页提供月历打卡（有训练的日子打点）、近 8 周训练趋势、PR 个人最好成绩榜（最大重量与最重单次容量），"
          "让进步「看得见」。首页则展示本周训练次数、时长、总容量与最近一次训练摘要。"),
    ("h3", "分享传播"),
    ("p", "训练完成后可生成分享卡：白卡简洁版式，可选相册照片作背景并自动叠加遮罩保证可读性，"
          "Canvas 本地绘制不依赖网络，可保存到相册发朋友圈，也可直接转发微信好友。"),
    ("h2", "3.4 交互设计"),
    ("bullet", "底部导航采用微信原生 tabBar（首页/训练/历史/我的），稳定且适配安全区"),
    ("bullet", "动作选择、模板选择使用半屏弹层，训练主流程不被打断"),
    ("bullet", "训练中返回有防误退确认，草稿自动保留"),
    ("bullet", "间歇倒计时到点震动提醒，训练节奏不靠眼睛盯表"),
    ("bullet", "录入错误即时提示（年龄/身高/体重/次数/重量范围校验，前后端一致）"),
    ("h2", "3.5 视觉设计：Notion 风格设计语言"),
    ("p", "整套 UI 基于 awesome-design-md 的 Notion 设计语言定制：深海军蓝 Hero（#0a1530）营造沉稳的健身工作台气质，"
          "紫色（#5645d4）作为唯一主行动色，粉彩卡片（tint-*）对应不同健身目标与标签，卡片以 1px 发丝线区分层级、"
          "不加投影，按钮采用 8px 圆角矩形而非胶囊。字体栈 Inter / PingFang SC / Microsoft YaHei，"
          "设计令牌（颜色、间距、圆角、字号）集中到 app.wxss，页面不硬编码。"),
    ("h2", "3.6 专业标准与数据依据"),
    ("p", "体质指标与推荐重量采用中国标准与国际通用公式，全部为服务端与前端共用的纯函数，保证结果一致："),
    ("table", {
        "headers": ["指标", "标准 / 公式", "说明"],
        "widths": [16, 40, 44],
        "rows": [
            ["BMI", "体重(kg) ÷ 身高(m)²", "<18.5 偏瘦；18.5–23.9 正常；24–27.9 超重；≥28 肥胖（中国标准）"],
            ["体脂参考", "中国常用健康参考值", "男性 <10% / 10–20% / 20–25% / >25%；女性 <15% / 15–25% / 25–30% / >30%"],
            ["基础代谢 BMR", "Mifflin-St Jeor", "男 = 10×体重 + 6.25×身高 − 5×年龄 + 5；女 = 10×体重 + 6.25×身高 − 5×年龄 − 161"],
            ["推荐重量区间", "体重 × 倍数", "每个动作按性别分别内置新手/中级/高级倍数区间，取 2.5kg 档；自重动作不给出重量"],
            ["动作库", "国际通用动作命名", "59 个动作中英文对照，按 7 大部位与器械（杠铃/哑铃/器械/绳索/自重）分类"],
        ],
    }),

    ("h1", "四、实践过程"),
    ("h2", "4.1 技术选型"),
    ("table", {
        "headers": ["层次", "选型", "理由"],
        "widths": [16, 30, 54],
        "rows": [
            ["客户端", "微信原生小程序（基础库 3.17.1）", "免安装、性能稳、与微信分享/相册能力天然集成"],
            ["后端", "微信云开发（云函数 + 云数据库）", "免运维、按量计费，openid 天然隔离用户数据"],
            ["逻辑层", "纯函数工具层（standards/stats/format）", "无 wx 依赖，可用 node:test 单元测试，前后端复用"],
            ["测试", "Node.js 24 + node:test + 开发者工具自动化", "TDD 先行，单测 + 模拟器全流程自动化回归"],
            ["工程化", "Git 分支管理 + 云开发 CLI", "版本可回溯，云函数一键部署"],
        ],
    }),
    ("h2", "4.2 总体架构"),
    ("diagram", "arch"),
    ("p", "前端页面通过统一工具层访问数据：utils/standards.js、stats.js、format.js 为纯函数逻辑；"
          "cloud.js 封装云函数调用并做目录缓存；storage.js 负责本地草稿与目录缓存兜底。"
          "云函数负责身份识别、数据校验、统计计算与种子数据下发，云数据库按 openid 隔离数据。"
          "即使云端不可用，本地种子数据与草稿机制也能保证「先记下来」，网络恢复后自动补传。"),
    ("h2", "4.3 数据模型设计"),
    ("table", {
        "headers": ["集合", "关键字段", "说明"],
        "widths": [18, 46, 36],
        "rows": [
            ["users", "openid、gender、age、heightCm、weightKg、goal、frequency、bmi、bmiLevel、bodyFatRef、bmr", "用户档案与体质指标，服务端计算回写"],
            ["workouts", "openid、date、startedAt、endedAt、durationSec、mode、templateId、templateName、exercises[]、totalSets、totalVolume、calories", "一次训练记录，exercises[] 内嵌动作与组数据"],
            ["exercises", "id、name、enName、bodyPart、equipment、weighted、pcts", "全局只读种子，59 个专业动作"],
            ["templates", "id、name、goal、genderHint、frequency、exercises[]", "全局只读种子，4 套训练模板"],
        ],
    }),
    ("p", "PR 个人最好成绩不单独建集合，由 stats 云函数从 workouts 实时现算，减少一个需要同步的集合，"
          "从根本上避免数据不一致。"),
    ("h2", "4.4 云函数实现"),
    ("table", {
        "headers": ["云函数", "职责", "超时"],
        "widths": [20, 62, 18],
        "rows": [
            ["login", "获取 openid，取或建用户记录，返回用户档案", "5s"],
            ["saveProfile", "校验资料，服务端计算 BMI/分级/体脂参考/BMR 并回写 users", "10s"],
            ["catalog", "并行返回动作库与模板种子数据", "10s"],
            ["saveWorkout", "校验训练记录（日期/模式/动作/热量），服务端计算总组数与总容量，写入 workouts", "10s"],
            ["stats", "按 openid 拉取最近 1000 条训练，现算本周概览、8 周趋势、PR、最近 5 条", "15s"],
            ["initDb", "幂等初始化：自动创建集合并同步种子数据（可重复执行）", "60s"],
        ],
    }),
    ("h2", "4.5 开发里程碑"),
    ("p", "项目经历了「初版上线 → 真机暴露问题 → 第一性原理重建 → 用户反馈迭代 → v1.2.0」的完整过程："),
    ("bullet", "2026-08-16：完成初版 v1.0.0 并上传，云函数与数据库初始化完成"),
    ("bullet", "2026-08-16：真机验收发现三类系统性问题——页面横向溢出、模板训练不可用、分享卡不可用，决定暂缓上线"),
    ("bullet", "2026-08-17：撰写第一性原理重建设计文档与实现计划，确定「最短闭环」产品原则"),
    ("bullet", "2026-08-17：按 TDD 完成重建——纯逻辑层 → 种子数据 → 5 个云函数 → 应用骨架 → 工具层 → 6 个页面 → 部署验收"),
    ("bullet", "2026-08-17~18：根据用户反馈迭代——动作库扩至 59 个、新增自定义动作、部位快捷筛选、动作间歇计时器、模板换入口、分享卡热量展示等"),
    ("bullet", "2026-08-18：15 个单元测试全绿、模拟器全流程自动化 0 报错，v1.2.0 上传微信后台（79.2 KB），进入提审准备"),
    ("h2", "4.6 测试与质量保障"),
    ("p", "采用 TDD：先写失败测试，再实现纯函数逻辑。测试覆盖 BMI 分级边界、男女体脂参考、BMR 公式、"
          "推荐重量区间、总组数/总容量、周概览/周趋势、PR 现算、格式化等，共 15 个用例全部通过。"
          "页面层通过开发者工具自动化脚本完成全流程回归：建档 → 首页 → 自由训练 → 模板训练 → 自定义动作 → "
          "部位筛选 → 换模板 → 热量与间歇计时 → 保存 → 历史 → 分享，均 0 报错。"),
    ("h2", "4.7 真机问题与修复"),
    ("table", {
        "headers": ["真机问题", "根因分析", "修复方案"],
        "widths": [26, 30, 44],
        "rows": [
            ["页面横向溢出，可左右拖动", "全局样式复用 + 内容盒模型导致宽度计算错误", "全局 box-sizing:border-box、页面 overflow-x:hidden、全部改用 rpx，390px 视口实测无溢出"],
            ["模板训练用不了", "模板强依赖云端加载，失败即中断", "模板并入 record 页半屏弹层，内置种子兜底，支持训练中换模板"],
            ["训练成果分享不了", "Canvas 参与布局宽度，绘制依赖网络与权限", "Canvas 离屏绘制（position:fixed 不占布局）、本地绘制、权限拒绝有明确提示与重试"],
            ["选动作只有胸，下拉看不到其他部位", "弹层滚动区高度依赖内容撑开，真机被压缩", "滚动区改为显式 60vh 高度，并新增部位快捷筛选标签（全部/胸/背/腿/臀腿/肩/手臂/核心）"],
            ["动作选择太少", "内置动作仅 12 个", "动作库扩充至 59 个（中英文术语 + 器械 + 推荐重量区间），并支持自定义动作"],
            ["无法记录间歇与热量", "功能缺失", "新增每组动作 1–30 分钟间歇倒计时（底部常驻 + 震动提醒）；热量字段保留供 Apple Watch 后续接入"],
        ],
    }),
    ("h2", "4.8 部署与上线"),
    ("p", "云函数已通过云开发 CLI 部署至云环境 cloudbase-d9gyqv3ea400083a0，数据库由 initDb 幂等初始化（users/workouts/exercises/templates 4 个集合），"
          "集合权限设为「仅创建者可读写」。小程序代码通过微信开发者工具上传，当前版本 v1.2.0（79.2 KB），"
          "类目为「工具 → 健康管理」，用户隐私保护指引已按实际调用情况填写（相册仅写入、选中的照片、明示用途）。"),

    ("h1", "五、功能实现详解"),
    ("h2", "5.1 首次建档引导（onboarding）"),
    ("bullet", "三步引导：隐私协议勾选 → 性别/年龄/身高/体重 → 健身目标与每周频率"),
    ("bullet", "输入范围校验（年龄 6–100、身高 80–250cm、体重 20–300kg），错误即时提示"),
    ("bullet", "提交后由云函数计算并展示 BMI、体脂参考区间、基础代谢 BMR，用户看到自己的「身体起点」"),
    ("h2", "5.2 首页（home）"),
    ("bullet", "深海军蓝 Hero 区：问候语 + 本周训练次数、时长、总容量三项指标"),
    ("bullet", "「开始训练」主按钮直达训练页；最近一次训练摘要卡"),
    ("bullet", "未建档自动跳转建档页"),
    ("h2", "5.3 训练记录页（record，核心页面）"),
    ("bullet", "自由训练 / 模板训练双模式；模板训练显示「当前模板」并可随时更换"),
    ("bullet", "59 个动作按 7 大部位分组，弹层顶部部位快捷标签，滚动区固定 60vh"),
    ("bullet", "每个动作展示器械类型与推荐重量区间（按性别 + 体重自动换算）"),
    ("bullet", "自定义动作：一键添加，本地存储，满足个性化需求"),
    ("bullet", "按组录入次数与重量，可加组/删组/删动作；自重动作重量为 0"),
    ("bullet", "每组动作可设 1–30 分钟间歇倒计时，底部常驻显示，到点震动提醒"),
    ("bullet", "底部常驻总组数/总容量，实时更新；「完成训练」校验后本地草稿 + 云端保存"),
    ("bullet", "训练中返回防误退，草稿自动保留，断网不丢数据"),
    ("h2", "5.4 历史页（history）"),
    ("bullet", "月历打卡：有训练的日子打点，可切换月份"),
    ("bullet", "近 8 周训练趋势图（按周统计次数与容量）"),
    ("bullet", "PR 个人最好成绩榜：每个动作的最大重量与最重单次容量及达成日期"),
    ("h2", "5.5 我的页（profile）"),
    ("bullet", "个人资料卡 + 体质指标卡（BMI/分级/体脂参考/BMR）"),
    ("bullet", "用户协议、隐私政策入口；注销账号入口（二次确认）"),
    ("bullet", "隐私合规：最小化收集、明示用途、联系方式可查"),
    ("h2", "5.6 分享卡（share）"),
    ("bullet", "白卡简洁版式：品牌名 + 「训练打卡」标签 + 总容量大字 + 动作摘要 + 日期"),
    ("bullet", "可选照片背景：相册选图自动叠加半透明深色遮罩并切换白色文字保证可读性"),
    ("bullet", "Canvas 本地绘制，不依赖网络；保存相册（发朋友圈）或直接转发好友"),
    ("bullet", "画布初始化失败/相册权限被拒时有明确提示，可重试或跳过"),

    ("h1", "六、技术亮点与创新"),
    ("bullet", "纯函数分层 + TDD：体质指标、统计、格式化均为无 wx 依赖的纯函数，前后端复用，15 个单测全绿"),
    ("bullet", "本地草稿兜底：保存先落本地再上云，云端失败保留草稿，启动自动补传，数据零丢失"),
    ("bullet", "PR 现算不落库：个人最好成绩由 stats 云函数实时计算，避免多集合同步带来的数据不一致"),
    ("bullet", "缓存版本号 + 24 小时过期：动作库更新可即时生效，避免旧缓存导致新动作不出现"),
    ("bullet", "服务端校验与统计：日期/模式/动作/热量/总组数/总容量均在云函数侧二次校验，前端防错、后端兜底"),
    ("bullet", "设计令牌 + 真机适配：全局 border-box、全 rpx、安全区、禁止横向滚动，解决真机适配的系统性问题"),
    ("bullet", "隐私与安全：云数据库「仅创建者可读写」、数据按 openid 隔离、最小化收集、注销入口"),

    ("h1", "七、测试与验收结果"),
    ("h2", "7.1 单元测试（15 个全部通过）"),
    ("table", {
        "headers": ["测试文件", "覆盖内容"],
        "widths": [30, 70],
        "rows": [
            ["standards.test.js", "BMI 计算与中国标准分级边界、男女体脂参考、BMR（Mifflin-St Jeor）、推荐重量区间、建档校验"],
            ["stats.test.js", "总组数/总容量、本周概览、8 周趋势、PR 现算（最大重量/最重单次容量与日期）"],
            ["format.test.js", "容量千分位与吨位、时长格式化、日期与星期格式化"],
        ],
    }),
    ("h2", "7.2 自动化 UI 全流程"),
    ("p", "使用微信开发者工具自动化脚本在模拟器上完成端到端回归：首次建档、首页概览、自由训练、模板训练、"
          "自定义动作、部位快捷筛选、换模板、热量字段与间歇倒计时、保存训练、历史月历/趋势/PR、分享卡生成，全部 0 报错；"
          "并单独校验 390px 视口下页面宽度无横向溢出。"),
    ("h2", "7.3 验收清单要点"),
    ("bullet", "首次建档：协议未勾选不可继续；非法输入有提示；指标结果与公式一致"),
    ("bullet", "训练记录：模板预置组数正确；组内重量/次数可编辑；总组数/总容量实时更新；未完成组有提示"),
    ("bullet", "历史统计：日历标记训练日、趋势正确、PR 列表正确"),
    ("bullet", "分享：白卡预览正确、照片背景可读、保存相册成功、转发好友成功"),
    ("bullet", "合规：登录前展示用户协议与隐私政策；数据库仅创建者可读写；提供注销入口"),

    ("h1", "八、部署与上线"),
    ("table", {
        "headers": ["项目", "信息"],
        "widths": [26, 74],
        "rows": [
            ["小程序名称", "牛来举铁"],
            ["AppID", "wx2a14c212978a5374"],
            ["云环境", "cloudbase-d9gyqv3ea400083a0"],
            ["云函数", "login / saveProfile / catalog / saveWorkout / stats / initDb（均已部署）"],
            ["数据库集合", "users / workouts / exercises / templates（已初始化，仅创建者可读写）"],
            ["当前版本", "v1.2.0（79.2 KB，已上传微信后台，待提交审核）"],
            ["服务类目", "工具 → 健康管理"],
            ["隐私保护", "隐私保护指引已配置：相册仅写入、选中的照片、明示用途；更新/生效日期 2026-08-17"],
        ],
    }),
    ("p", "上线链路：开发者工具上传 → 公众平台提交审核 → 审核通过后发布。云开发自带域名与存储，无需服务器与备案。"
          "上线后小程序码云函数可恢复使用，分享卡可进一步升级为带小程序码的版本。"),

    ("h1", "九、价值与展望"),
    ("h2", "9.1 社会价值"),
    ("p", "在「健康中国」背景下，产品用极低门槛帮助用户建立「记录 → 反馈 → 坚持」的正循环：专业标准让新手少走弯路，"
          "数据反馈让坚持可见，分享卡把自律扩散为社交影响力，带动更多人科学健身。"),
    ("h2", "9.2 未来规划"),
    ("bullet", "Apple Watch / 健康数据接入：训练热量自动同步（数据字段已预留），形成更完整的健康画像"),
    ("bullet", "分享卡升级：叠加小程序码，形成传播裂变"),
    ("bullet", "个性化训练计划：基于档案、目标与历史数据自动生成周期化训练计划"),
    ("bullet", "动作教学：为 59 个动作补充图文/视频标准动作演示"),
    ("bullet", "社交激励：好友 PK、训练周报、连续打卡成就体系"),
    ("bullet", "多端与数据导出：支持数据导出与更多穿戴设备接入"),

    ("pagebreak", None),
    ("h1", "附录 A：动作库清单（59 个）"),
    ("table", {
        "headers": ["部位", "动作（中文）", "英文名", "器械"],
        "widths": [10, 30, 42, 18],
        "rows": [
            ["胸", "杠铃卧推", "Barbell Bench Press", "杠铃"],
            ["胸", "上斜杠铃卧推", "Incline Barbell Press", "杠铃"],
            ["胸", "下斜杠铃卧推", "Decline Barbell Press", "杠铃"],
            ["胸", "哑铃卧推", "Dumbbell Bench Press", "哑铃"],
            ["胸", "上斜哑铃卧推", "Incline Dumbbell Press", "哑铃"],
            ["胸", "坐姿推胸", "Machine Chest Press", "器械"],
            ["胸", "蝴蝶机夹胸", "Pec Deck Fly", "器械"],
            ["胸", "绳索夹胸", "Cable Crossover", "绳索"],
            ["胸", "哑铃飞鸟", "Dumbbell Fly", "哑铃"],
            ["胸", "俯卧撑", "Push-Up", "自重"],
            ["胸", "双杠臂屈伸", "Chest Dip", "自重"],
            ["背", "硬拉", "Deadlift", "杠铃"],
            ["背", "坐姿划船", "Seated Cable Row", "器械"],
            ["背", "杠铃俯身划船", "Barbell Bent-Over Row", "杠铃"],
            ["背", "单臂哑铃划船", "One-Arm Dumbbell Row", "哑铃"],
            ["背", "高位下拉", "Lat Pulldown", "绳索"],
            ["背", "引体向上", "Pull-Up", "自重"],
            ["背", "反手引体向上", "Chin-Up", "自重"],
            ["背", "T杠划船", "T-Bar Row", "杠铃"],
            ["背", "杠铃耸肩", "Barbell Shrug", "杠铃"],
            ["背", "面拉", "Face Pull", "绳索"],
            ["腿", "杠铃深蹲", "Barbell Back Squat", "杠铃"],
            ["腿", "杠铃前蹲", "Front Squat", "杠铃"],
            ["腿", "腿举", "Leg Press", "器械"],
            ["腿", "哈克深蹲", "Hack Squat", "器械"],
            ["腿", "高脚杯深蹲", "Goblet Squat", "哑铃"],
            ["腿", "罗马尼亚硬拉", "Romanian Deadlift", "杠铃"],
            ["腿", "哑铃弓步蹲", "Dumbbell Lunge", "哑铃"],
            ["腿", "腿屈伸", "Leg Extension", "器械"],
            ["腿", "俯卧腿弯举", "Lying Leg Curl", "器械"],
            ["腿", "站姿提踵", "Standing Calf Raise", "器械"],
            ["腿", "相扑硬拉", "Sumo Deadlift", "杠铃"],
            ["臀腿", "臀桥", "Hip Bridge", "自重"],
            ["臀腿", "杠铃臀推", "Barbell Hip Thrust", "杠铃"],
            ["臀腿", "绳索后踢腿", "Cable Glute Kickback", "绳索"],
            ["臀腿", "坐姿腿外展", "Hip Abduction", "器械"],
            ["肩", "哑铃推举", "Dumbbell Shoulder Press", "哑铃"],
            ["肩", "杠铃推举", "Overhead Press", "杠铃"],
            ["肩", "阿诺德推举", "Arnold Press", "哑铃"],
            ["肩", "哑铃侧平举", "Dumbbell Lateral Raise", "哑铃"],
            ["肩", "哑铃前平举", "Dumbbell Front Raise", "哑铃"],
            ["肩", "俯身哑铃飞鸟", "Rear Delt Fly", "哑铃"],
            ["肩", "绳索面拉变式", "Cable Rear Delt Fly", "绳索"],
            ["手臂", "杠铃弯举", "Barbell Curl", "杠铃"],
            ["手臂", "哑铃弯举", "Dumbbell Curl", "哑铃"],
            ["手臂", "锤式弯举", "Hammer Curl", "哑铃"],
            ["手臂", "绳索下压", "Triceps Pushdown", "绳索"],
            ["手臂", "仰卧臂屈伸", "Skull Crusher", "杠铃"],
            ["手臂", "窄距卧推", "Close-Grip Bench Press", "杠铃"],
            ["手臂", "绳索弯举", "Cable Curl", "绳索"],
            ["手臂", "双杠臂屈伸（手臂）", "Triceps Dip", "自重"],
            ["核心", "卷腹", "Crunch", "自重"],
            ["核心", "平板支撑", "Plank", "自重"],
            ["核心", "悬垂举腿", "Hanging Leg Raise", "自重"],
            ["核心", "俄罗斯转体", "Russian Twist", "自重"],
            ["核心", "绳索卷腹", "Cable Crunch", "绳索"],
            ["核心", "健腹轮", "Ab Wheel Rollout", "自重"],
            ["核心", "侧平板支撑", "Side Plank", "自重"],
            ["核心", "仰卧抬腿", "Lying Leg Raise", "自重"],
        ],
    }),
    ("h1", "附录 B：训练模板清单（4 套）"),
    ("table", {
        "headers": ["模板", "目标人群", "特点"],
        "widths": [26, 26, 48],
        "rows": [
            ["全身分化 · 男版", "男性新手", "每周 3 次全身训练，覆盖胸背腿肩核心"],
            ["全身分化 · 女版", "女性新手", "按女性推荐重量倍数设计，兼顾臀腿塑形"],
            ["推-拉-腿 PPL", "进阶爱好者", "经典三分化，训练频率高、容量大"],
            ["上下肢分化", "进阶爱好者", "四天上下肢分化，力量与容量兼顾"],
        ],
    }),
    ("h1", "附录 C：版本迭代记录"),
    ("table", {
        "headers": ["版本/提交", "内容"],
        "widths": [24, 76],
        "rows": [
            ["v1.0.0", "初版上线，云函数与数据库初始化完成；真机验收暴露横向溢出/模板/分享三类问题"],
            ["v2 设计", "第一性原理重建设计文档与实现计划，确立「最短闭环」原则"],
            ["重建版", "纯逻辑层 TDD、5 个云函数、应用骨架、6 个页面、本地兜底与分享卡全部重写"],
            ["迭代 1", "动作库 12 → 59 个（含中英文术语），新增自定义动作与云端种子同步"],
            ["迭代 2", "动作选择弹层滚动修复 + 部位快捷筛选标签"],
            ["迭代 3", "分享卡以热量替代总容量；训练记录热量；动作间歇倒计时 1–30 分钟（底部常驻 + 震动）"],
            ["迭代 4", "模板训练顶部显示当前模板 + 换模板入口"],
            ["迭代 5", "去掉训练记录页热量输入框（保留数据字段，供后续 Apple Watch 接入）"],
            ["v1.2.0", "15 单测全绿 + 自动化全流程 0 报错，上传微信后台，待提交审核"],
        ],
    }),
]


# ---------- PDF helpers ----------
class Doc(BaseDocTemplate):
    def afterFlowable(self, flowable):
        if isinstance(flowable, Paragraph):
            style = flowable.style.name
            text = flowable.getPlainText()
            if text == "目录":
                return
            if style == "H1":
                key = "h1-" + hashlib.md5(text.encode("utf-8")).hexdigest()[:10]
                self.canv.bookmarkPage(key)
                self.notify("TOCEntry", (0, text, self.page, key))
            elif style == "H2":
                key = "h2-" + hashlib.md5(text.encode("utf-8")).hexdigest()[:10]
                self.canv.bookmarkPage(key)
                self.notify("TOCEntry", (1, text, self.page, key))


def cover_page(canv, doc):
    canv.saveState()
    w, h = A4
    canv.setFillColor(NAVY)
    canv.rect(0, 0, w, h, stroke=0, fill=1)
    canv.setFillColor(PURPLE)
    canv.rect(0, h - 8 * mm, w, 8 * mm, stroke=0, fill=1)
    canv.setFillColor(colors.white)
    canv.setFont(FONT_B, 34)
    canv.drawCentredString(w / 2, h - 95 * mm, "牛来举铁")
    canv.setFont(FONT, 15)
    canv.setFillColor(colors.HexColor("#c8c4be"))
    canv.drawCentredString(w / 2, h - 103 * mm, "微信健身记录小程序 · 赛事说明文档")
    canv.setFont(FONT, 10.5)
    canv.setFillColor(colors.HexColor("#a4a097"))
    canv.drawCentredString(w / 2, h - 128 * mm, "记一次训练  →  看见进步  →  分享成果")
    canv.setStrokeColor(colors.HexColor("#3a3f55"))
    canv.setLineWidth(0.8)
    canv.line(45 * mm, h - 138 * mm, w - 45 * mm, h - 138 * mm)
    canv.setFillColor(colors.HexColor("#a4a097"))
    canv.setFont(FONT, 9.5)
    meta = ["版本：v1.2.0    开发方式：微信原生小程序 + 微信云开发", "AppID：wx2a14c212978a5374    云环境：cloudbase-d9gyqv3ea400083a0", "2026 年 8 月"]
    y = h - 148 * mm
    for line in meta:
        canv.drawCentredString(w / 2, y, line)
        y -= 7 * mm
    canv.restoreState()


def body_page(canv, doc):
    canv.saveState()
    w, h = A4
    canv.setFillColor(colors.HexColor("#a4a097"))
    canv.setFont(FONT, 8.5)
    canv.drawString(20 * mm, 12 * mm, "牛来举铁 · 微信小程序赛事说明文档")
    canv.drawRightString(w - 20 * mm, 12 * mm, "第 %d 页" % doc.page)
    canv.setStrokeColor(HAIR)
    canv.setLineWidth(0.5)
    canv.line(20 * mm, 15 * mm, w - 20 * mm, 15 * mm)
    canv.restoreState()


def make_diagram(kind):
    d = Drawing(170 * mm, 58 * mm)
    if kind == "loop":
        items = ["首次建档", "记录训练", "查看进步", "分享成果"]
        n = len(items)
        box_w, box_h = 34 * mm, 18 * mm
        total = n * box_w + (n - 1) * 12 * mm
        x0 = (170 * mm - total) / 2
        y = 22 * mm
        for i, t in enumerate(items):
            x = x0 + i * (box_w + 12 * mm)
            d.add(Rect(x, y, box_w, box_h, fillColor=SURFACE, strokeColor=PURPLE, strokeWidth=1.2))
            d.add(String(x + box_w / 2, y + box_h / 2 - 3, t, fontName=FONT, fontSize=11,
                         fillColor=NAVY, textAnchor="middle"))
            if i < n - 1:
                ax = x + box_w + 1.2 * mm
                ay = y + box_h / 2
                d.add(Line(ax, ay, ax + 9.6 * mm, ay, strokeColor=PURPLE, strokeWidth=1.2))
                d.add(Polygon([ax + 10.4 * mm, ay, ax + 8.4 * mm, ay - 1.6 * mm,
                               ax + 8.4 * mm, ay + 1.6 * mm], fillColor=PURPLE, strokeColor=None))
        d.add(String(85 * mm, 6 * mm, "最短核心闭环", fontName=FONT, fontSize=9,
                     fillColor=SLATE, textAnchor="middle"))
    else:
        layers = [
            ("小程序页面层", "onboarding / home / record / history / profile / share（6 页面，原生 tabBar）"),
            ("前端工具层", "standards（体质指标） / stats（统计） / format（格式化） / cloud（云调用） / storage（草稿+缓存）"),
            ("云函数层", "login / saveProfile / catalog / saveWorkout / stats / initDb"),
            ("云数据库层", "users / workouts / exercises / templates（仅创建者可读写）"),
        ]
        n = len(layers)
        bh = 10.5 * mm
        gap = 5 * mm
        total_h = n * bh + (n - 1) * gap
        y = (58 * mm - total_h) / 2 + 1 * mm
        for i, (title, desc) in enumerate(layers):
            yy = y + (n - 1 - i) * (bh + gap)
            d.add(Rect(8 * mm, yy, 154 * mm, bh, fillColor=colors.white, strokeColor=PURPLE, strokeWidth=1.1))
            d.add(String(13 * mm, yy + bh / 2 - 3, title, fontName=FONT_B, fontSize=9.5, fillColor=NAVY))
            d.add(String(43 * mm, yy + bh / 2 - 3, desc, fontName=FONT, fontSize=8.4, fillColor=SLATE))
            if i < n - 1:
                ax = 85 * mm
                ay = yy - gap + 1.4 * mm
                d.add(Line(ax, ay + 1.2 * mm, ax, ay + gap - 1.2 * mm, strokeColor=PURPLE, strokeWidth=1))
                d.add(Polygon([ax, ay, ax - 1.6 * mm, ay + 1.8 * mm, ax + 1.6 * mm, ay + 1.8 * mm],
                              fillColor=PURPLE, strokeColor=None))
    return d


def render_pdf():
    frame_body = Frame(20 * mm, 18 * mm, A4[0] - 40 * mm, A4[1] - 42 * mm, id="body")
    frame_cover = Frame(0, 0, A4[0], A4[1], id="cover")
    doc = Doc(PDF_PATH, pagesize=A4, title="牛来举铁 - 微信小程序赛事说明文档",
              author="牛来举铁 开发团队", subject="创作背景 / 创作思路 / 实践过程")
    doc.addPageTemplates([
        PageTemplate(id="cover", frames=[frame_cover], onPage=cover_page),
        PageTemplate(id="body", frames=[frame_body], onPage=body_page),
    ])

    story = []
    story.append(NextPageTemplate("body"))
    story.append(PageBreak())

    toc = TableOfContents()
    toc.levelStyles = [S_TOC1, S_TOC2]
    story.append(Paragraph("目录", S_H1))
    story.append(toc)
    story.append(Spacer(1, 8 * mm))

    for kind, payload in CONTENT:
        if kind == "h1":
            story.append(Paragraph(payload, S_H1))
        elif kind == "h2":
            story.append(Paragraph(payload, S_H2))
        elif kind == "h3":
            story.append(Paragraph(payload, S_H3))
        elif kind == "p":
            story.append(Paragraph(payload, S_P))
        elif kind == "bullet":
            story.append(Paragraph(payload, S_BULLET, bulletText="•"))
        elif kind == "diagram":
            story.append(make_diagram(payload))
            story.append(Spacer(1, 4 * mm))
        elif kind == "table":
            spec = payload
            header = [Paragraph(h, S_CELLB) for h in spec["headers"]]
            rows = [[Paragraph(c, S_CELL) for c in row] for row in spec["rows"]]
            data = [header] + rows
            widths = [w * mm for w in spec["widths"]]
            tbl = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), SURFACE),
                ("GRID", (0, 0), (-1, -1), 0.6, HAIR),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 3.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fbfaf9")]),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 4 * mm))
        elif kind == "pagebreak":
            story.append(PageBreak())

    doc.multiBuild(story)
    return PDF_PATH


# ---------- DOCX ----------
def render_docx():
    import docx
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    def set_ea(style, name="微软雅黑"):
        style.font.name = name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), name)

    h1 = doc.styles["Heading 1"]
    set_ea(h1)
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x0A, 0x15, 0x30)
    h2 = doc.styles["Heading 2"]
    set_ea(h2)
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x56, 0x45, 0xD4)
    h3 = doc.styles["Heading 3"]
    set_ea(h3)
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Cover
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("牛来举铁")
    r.font.size = Pt(34)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x0A, 0x15, 0x30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("微信健身记录小程序 · 赛事说明文档")
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(0x5D, 0x5B, 0x54)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("记一次训练  →  看见进步  →  分享成果")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xA4, 0xA0, 0x97)
    doc.add_paragraph()
    for line in ["版本：v1.2.0    开发方式：微信原生小程序 + 微信云开发",
                 "AppID：wx2a14c212978a5374    云环境：cloudbase-d9gyqv3ea400083a0",
                 "2026 年 8 月"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xA4, 0xA0, 0x97)
    doc.add_page_break()

    # TOC placeholder
    p = doc.add_paragraph()
    r = p.add_run("目录")
    r.font.size = Pt(16)
    r.font.bold = True
    toc_items = [c[1] for c in CONTENT if c[0] == "h1"]
    for item in toc_items:
        doc.add_paragraph(item, style="Heading 3")
    doc.add_page_break()

    def add_table(headers, rows, widths):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = ""
            run = hdr[i].paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x0A, 0x15, 0x30)
            hdr[i].paragraphs[0].paragraph_format.space_before = Pt(2)
            hdr[i].paragraphs[0].paragraph_format.space_after = Pt(2)
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                run = cells[i].paragraphs[0].add_run(str(val))
                run.font.size = Pt(9)
                cells[i].paragraphs[0].paragraph_format.space_before = Pt(2)
                cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)
        doc.add_paragraph()

    for kind, payload in CONTENT:
        if kind == "h1":
            doc.add_heading(payload, level=1)
        elif kind == "h2":
            doc.add_heading(payload, level=2)
        elif kind == "h3":
            doc.add_heading(payload, level=3)
        elif kind == "p":
            doc.add_paragraph(payload)
        elif kind == "bullet":
            doc.add_paragraph(payload, style="List Bullet")
        elif kind == "diagram":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("（架构/闭环示意图见 PDF 版对应章节）")
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xA4, 0xA0, 0x97)
        elif kind == "table":
            spec = payload
            add_table(spec["headers"], spec["rows"], spec["widths"])
        elif kind == "pagebreak":
            doc.add_page_break()

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(render_pdf())
    print(render_docx())
